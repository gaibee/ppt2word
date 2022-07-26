
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Mm
import json
import os, sys
from tkinter import filedialog
from components import ppt, border
import unicodedata





def remove_control_characters(s):
    s = s.encode('utf-8').decode('utf-8', 'ignore')
    return "".join(ch if unicodedata.category(ch) not in ['Cc', 'Cf', 'Cs', 'Co', 'Cn'] else '\n' for ch in s )

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


# Working dir을 현재 파이썬 파일의 dir로 설정
current_path = os.path.dirname(os.path.abspath(__file__))
os.chdir(current_path)


# 설정 불러오기
config = {}
try:
    with open('config.json', 'r', encoding='utf-8') as f:
        config = json.load(f)
except:
    print('config.json 파일을 찾을 수 없습니다.')
    sys.exit()

# 실행 변수 불러오기
if len(sys.argv) == 1:

    # 실행인자가 없으면 직접 ppt 파일을 선택하게 한다.
    file_path = filedialog.askopenfilename(initialdir=os.getcwd(), title = "ppt 파일을 선택 해 주세요", filetypes = (("*.pptx","*pptx"),("*.ppt","*ppt")))
    if file_path == '':
        sys.exit()

    print(f'선택된 ppt 파일 : {file_path}')

    # ppt 파일 -> png로 변환
    if not os.path.exists('png'):
        os.makedirs('png')
        print('png 폴더 생성')

    png_full_path = os.path.join( os.getcwd(), 'png' )
    print(f'{file_path} ppt 파일을 {png_full_path} 폴더에 png 파일로 저장중...')
    ppt.slide_to_image(png_full_path, file_path)
    print('- 완료')
    
else:
    file_path = sys.argv[1]

print('슬라이드 노트 불러오는 중...')
notes = ppt.get_slides_note(file_path)
print('- 완료')

print(notes)

print('Docx 파일 생성중...')
doc = Document()
table = doc.add_table(rows=len(notes), cols=2)
for i, row in enumerate(table.rows):
    print(f'슬라이드{i+1} 변환중...')
    
    cells = row.cells
    
    # 너비 높이 설정
    row.height = Cm(config['ROW_HEIGHT'])
    cells[0].width = Cm(config['COLUMN_PPT_WIDTH'])
    cells[1].width = Cm(config['COLUMN_NOTE_WIDTH'])
    
    # 야매로 보더 추가 (각 페이지의 맨 하단의 보더는 그리지 말아야함)
    if i % 3 != 2:
        border.set_cell_border(cells[0], bottom={"sz": 3*8, "val": "single", "color": "#5B9BD5", "space": "0"})
        border.set_cell_border(cells[1], bottom={"sz": 3*8, "val": "single", "color": "#5B9BD5", "space": "0"})
    
    # PPT 그림 추가
    # SIZE_MODE가 width이면 height는 그림 비율로 자동조정
    run = cells[0].paragraphs[0].add_run()
    if config['SIZE_MODE']=='width':
        run.add_picture(f'png/슬라이드{i+1}.PNG', width=Cm(config['PPT_WIDTH']))
    elif config['SIZE_MODE']=='height':
        run.add_picture(f'png/슬라이드{i+1}.PNG', height=Cm(config['PPT_HEIGHT']))
    
    # Note 텍스트 추가
    paragraph = cells[1].paragraphs[0]
    for text in remove_control_characters(notes[i]).split('\n'):
        
        # 단락 간격 설정
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1

        run = paragraph.add_run()
        
        # 특수문자 및 유니코드 이상한거 지우기
        run.text = remove_control_characters(text)
        
        # 폰트랑 글자 크기 설정
        run.font.name = config['FONT']
        run.font.size = Pt(config['FONT_SIZE'])
        run._element.rPr.rFonts.set(qn('w:eastAsia'), config['FONT'])
        
        paragraph = cells[1].add_paragraph()
        
    delete_paragraph(paragraph)

    # 가로 가운데 정렬
    cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for paragraph in cells[1].paragraphs:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 세로 가운데 정렬
    cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
# 여백 및 용지 크기 설정
for section in doc.sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.page_height = Mm(297)
    section.page_width = Mm(210)

doc.save('save.docx')
print('save.docx로 저장 완료')